#!/usr/bin/env python3
"""Render one Markdown source to BOTH .docx and .pdf.

Single source of truth on purpose: the two files are produced from the same parsed block list, so
they cannot drift. Supports headings, paragraphs, bullet/numbered lists, GitHub-style tables,
fenced code, blockquotes, and inline **bold** / *italic* / `code`.
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                                KeepTogether, PageBreak, Image as RLImage)
from reportlab.lib.utils import ImageReader

# ---------- parse ----------
def parse(md):
    blocks, lines, i = [], md.split("\n"), 0
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith("```"):
            buf = []; i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1; blocks.append(("code", buf)); continue
        mi = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", ln.strip())
        if mi:
            blocks.append(("img", (mi.group(2), mi.group(1)))); i += 1; continue
        m = re.match(r"^(#{1,4})\s+(.*)", ln)
        if m:
            blocks.append(("h%d" % len(m.group(1)), m.group(2).strip())); i += 1; continue
        if ln.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i+1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:|-]+$", "|".join(cells)):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows)); continue
        if re.match(r"^\s*[-*+]\s+", ln):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*+]\s+", "", lines[i])); i += 1
            blocks.append(("ul", items)); continue
        if re.match(r"^\s*\d+[.)]\s+", ln):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+[.)]\s+", "", lines[i])); i += 1
            blocks.append(("ol", items)); continue
        if ln.strip().startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip()); i += 1
            blocks.append(("quote", " ".join(buf))); continue
        if ln.strip() == "" or set(ln.strip()) <= {"-", "*", "_"} and len(ln.strip()) >= 3:
            i += 1; continue
        buf = []
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,4}\s|\s*[-*+]\s|\s*\d+[.)]\s|\||>|```)", lines[i]):
            buf.append(lines[i].strip()); i += 1
        if buf: blocks.append(("p", " ".join(buf)))
    return blocks

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[.+?\]\(.+?\))")
def strip_md(t):
    t = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", t)
    return t.replace("**", "").replace("`", "").replace("*", "")

# ---------- docx ----------
def add_runs(p, text):
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Menlo"; r.font.size = Pt(9)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            p.add_run(tok[1:-1]).italic = True
        elif tok.startswith("["):
            m = re.match(r"\[(.+?)\]\(.+?\)", tok); p.add_run(m.group(1) if m else tok)
        else:
            p.add_run(tok)

def to_docx(blocks, out):
    d = Document()
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
    for sec in d.sections:
        sec.left_margin = sec.right_margin = Inches(1.0)
        sec.top_margin = sec.bottom_margin = Inches(0.9)
    for kind, val in blocks:
        if kind == "h1":
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(strip_md(val)); r.bold = True; r.font.size = Pt(17)
            p.paragraph_format.space_after = Pt(12)
        elif kind in ("h2", "h3", "h4"):
            sz = {"h2": 13, "h3": 11.5, "h4": 10.5}[kind]
            p = d.add_paragraph(); r = p.add_run(strip_md(val)); r.bold = True; r.font.size = Pt(sz)
            p.paragraph_format.space_before = Pt(12 if kind == "h2" else 8)
            p.paragraph_format.space_after = Pt(4)
        elif kind == "p":
            add_runs(d.add_paragraph(), val)
        elif kind == "ul":
            for it in val: add_runs(d.add_paragraph(style="List Bullet"), it)
        elif kind == "ol":
            for it in val: add_runs(d.add_paragraph(style="List Number"), it)
        elif kind == "quote":
            p = d.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35)
            add_runs(p, val)
            for r in p.runs: r.italic = True
        elif kind == "code":
            for ln in val:
                p = d.add_paragraph(); r = p.add_run(ln or " ")
                r.font.name = "Menlo"; r.font.size = Pt(8.5)
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.left_indent = Inches(0.25)
        elif kind == "img":
            path, cap = val
            if Path(path).exists():
                p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(6.3))
                if cap:
                    c = d.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r = c.add_run(strip_md(cap)); r.font.size = Pt(8.5); r.italic = True
                    c.paragraph_format.space_after = Pt(10)
        elif kind == "table":
            hdr, body = val[0], val[1:]
            t = d.add_table(rows=1, cols=len(hdr)); t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, c in enumerate(hdr):
                cell = t.rows[0].cells[j]; cell.text = ""
                pr = cell.paragraphs[0]; r = pr.add_run(strip_md(c)); r.bold = True; r.font.size = Pt(9)
            for row in body:
                cells = t.add_row().cells
                for j, c in enumerate(row[:len(hdr)]):
                    cells[j].text = ""
                    pr = cells[j].paragraphs[0]; r = pr.add_run(strip_md(c)); r.font.size = Pt(9)
            d.add_paragraph().paragraph_format.space_after = Pt(4)
    d.save(out)

# ---------- pdf ----------
def to_pdf(blocks, out):
    ss = getSampleStyleSheet()
    S = {
        "h1": ParagraphStyle("h1", parent=ss["Title"], fontSize=17, leading=21, spaceAfter=14),
        "h2": ParagraphStyle("h2", parent=ss["Heading1"], fontSize=13, leading=16, spaceBefore=14, spaceAfter=5),
        "h3": ParagraphStyle("h3", parent=ss["Heading2"], fontSize=11.5, leading=14, spaceBefore=10, spaceAfter=4),
        "h4": ParagraphStyle("h4", parent=ss["Heading3"], fontSize=10.5, leading=13, spaceBefore=8, spaceAfter=3),
        "p":  ParagraphStyle("p", parent=ss["BodyText"], fontSize=9.8, leading=13.5, spaceAfter=6, alignment=4),
        "li": ParagraphStyle("li", parent=ss["BodyText"], fontSize=9.8, leading=13.5, leftIndent=16, spaceAfter=3),
        "q":  ParagraphStyle("q", parent=ss["BodyText"], fontSize=9.8, leading=13.5, leftIndent=22,
                             textColor=colors.HexColor("#444444"), spaceAfter=6),
        "code": ParagraphStyle("code", parent=ss["BodyText"], fontName="Courier", fontSize=8,
                               leading=9.6, leftIndent=14, spaceAfter=0),
        "cell": ParagraphStyle("cell", parent=ss["BodyText"], fontSize=8, leading=10, spaceAfter=0),
        "cellb": ParagraphStyle("cellb", parent=ss["BodyText"], fontSize=8, leading=10, spaceAfter=0,
                                fontName="Helvetica-Bold"),
    }
    def rt(t):
        t = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", t)
        t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"`(.+?)`", r'<font face="Courier" size="8.5">\1</font>', t)
        t = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", t)
        return t
    story, W = [], LETTER[0] - 2*inch
    for kind, val in blocks:
        if kind in ("h1", "h2", "h3", "h4"):
            story.append(Paragraph(rt(val), S[kind]))
        elif kind == "p":
            story.append(Paragraph(rt(val), S["p"]))
        elif kind == "ul":
            for it in val: story.append(Paragraph("• " + rt(it), S["li"]))
            story.append(Spacer(1, 4))
        elif kind == "ol":
            for n, it in enumerate(val, 1): story.append(Paragraph(f"{n}. " + rt(it), S["li"]))
            story.append(Spacer(1, 4))
        elif kind == "quote":
            story.append(Paragraph(rt(val), S["q"]))
        elif kind == "code":
            for ln in val: story.append(Paragraph(rt(ln) or "&nbsp;", S["code"]))
            story.append(Spacer(1, 6))
        elif kind == "img":
            path, cap = val
            if Path(path).exists():
                ir = ImageReader(path); iw, ih = ir.getSize()
                w = min(W, 6.3*inch); h = w*ih/iw
                grp = [RLImage(path, width=w, height=h)]
                if cap:
                    grp += [Spacer(1, 3), Paragraph("<i>"+rt(cap)+"</i>",
                            ParagraphStyle("cap", parent=S["p"], fontSize=8.2, leading=10.5, alignment=0))]
                story.append(KeepTogether(grp)); story.append(Spacer(1, 10))
        elif kind == "table":
            hdr, body = val[0], val[1:]
            n = len(hdr)
            data = [[Paragraph(rt(c), S["cellb"]) for c in hdr]]
            for row in body:
                row = (row + [""]*n)[:n]
                data.append([Paragraph(rt(c), S["cell"]) for c in row])
            first = max(1.0, W/n/inch*1.6) if n <= 4 else W/n/inch*1.25
            widths = [min(first, W/inch*0.34)*inch] + [(W - min(first, W/inch*0.34)*inch)/(n-1)]*(n-1) if n > 1 else [W]
            t = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF7")),
                ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#9BB0CC")),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("TOPPADDING", (0,0), (-1,-1), 3), ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F6F8FB")]),
            ]))
            story.append(t); story.append(Spacer(1, 9))
    SimpleDocTemplate(str(out), pagesize=LETTER,
                      leftMargin=inch, rightMargin=inch, topMargin=0.9*inch, bottomMargin=0.9*inch,
                      title="PEEK-Bench").build(story)

if __name__ == "__main__":
    src = Path(sys.argv[1]); stem = src.with_suffix("")
    blocks = parse(src.read_text())
    to_docx(blocks, str(stem) + ".docx")
    to_pdf(blocks, str(stem) + ".pdf")
    nt = sum(1 for k, _ in blocks if k == "table")
    print(f"  parsed {len(blocks)} blocks ({nt} tables)")
    for e in (".docx", ".pdf"):
        p = Path(str(stem) + e); print(f"  {p.name}: {p.stat().st_size/1024:.0f} KB")
